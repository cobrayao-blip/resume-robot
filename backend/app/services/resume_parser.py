import pdfplumber
import docx
import logging
from typing import Dict, Any
from .deepseek_service import deepseek_service

logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self, deepseek_service_instance=None):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
        # 如果提供了deepseek_service实例，使用它；否则使用全局实例
        self.deepseek_service = deepseek_service_instance or deepseek_service

    async def parse_resume_file(self, file_path: str, file_type: str, user=None, db_session=None, deepseek_service_instance=None) -> Dict[str, Any]:
        """
        解析简历文件并返回结构化数据
        Args:
            file_path: 文件路径
            file_type: 文件类型（pdf/docx）
            user: 用户对象（用于动态选择LLM配置）
            db_session: 数据库会话（用于动态选择LLM配置）
            deepseek_service_instance: LLM服务实例（向后兼容，已废弃）
        """
        import time
        total_start = time.time()
        
        try:
            # 提取文本内容
            extract_start = time.time()
            raw_text = self._extract_text(file_path, file_type)
            extract_elapsed = time.time() - extract_start
            text_length = len(raw_text)
            
            logger.info(f"[文件解析] 文本提取完成: 耗时{extract_elapsed:.2f}秒, 文本长度: {text_length}字符")
            
            # 放宽文本长度阈值，如果文本很短，记录警告但继续尝试解析
            if not raw_text or len(raw_text.strip()) < 10:
                logger.warning(f"[文件解析] 提取的文本很短: {len(raw_text.strip()) if raw_text else 0} 字符，但继续尝试AI解析")
                # 如果文本为空，使用默认提示文本
                if not raw_text or not raw_text.strip():
                    raise ValueError("无法从文件中提取有效文本内容，请检查文件格式是否正确")
            
            # 预处理文本：清理和优化，减少AI处理负担
            preprocess_start = time.time()
            cleaned_text = self._preprocess_text(raw_text)
            preprocess_elapsed = time.time() - preprocess_start
            cleaned_length = len(cleaned_text)
            reduction = text_length - cleaned_length
            logger.info(
                f"[文本预处理] 完成: 耗时{preprocess_elapsed:.2f}秒, "
                f"原始长度: {text_length}字符, 清理后: {cleaned_length}字符, "
                f"减少: {reduction}字符 ({reduction*100/text_length:.1f}%), "
                f"是否截断: {'是' if '[文本已截断' in cleaned_text or '[注意：文本已截断' in cleaned_text else '否'}"
            )
            
            # 使用 LLM 解析文本（使用清理后的文本）
            ai_start = time.time()
            # 使用传入的deepseek_service实例，如果没有则使用self.deepseek_service
            service_to_use = deepseek_service_instance or self.deepseek_service
            # 使用精简基础模型 schema 的 V2 解析方法
            structured_data = await service_to_use.parse_resume_text_v2(
                cleaned_text,
                user,
                db_session
            )
            structured_data = self._apply_dynamic_field_cleanup(structured_data)
            # 将增强格式转换为兼容格式，同时保留增强信息
            structured_data = self._normalize_enhanced_data(structured_data)
            ai_elapsed = time.time() - ai_start
            
            total_elapsed = time.time() - total_start
            work_count = len(structured_data.get('work_experiences', []))
            edu_count = len(structured_data.get('education', []))
            
            logger.info(
                f"[文件解析] 解析成功: 总耗时{total_elapsed:.2f}秒 "
                f"(文本提取: {extract_elapsed:.2f}秒, 预处理: {preprocess_elapsed:.2f}秒, AI解析: {ai_elapsed:.2f}秒), "
                f"提取信息: {work_count}个工作经历, {edu_count}个教育背景"
            )
            
            return structured_data
            
        except Exception as e:
            logger.error(f"简历文件解析失败: {e}")
            raise

    def _extract_text(self, file_path: str, file_type: str) -> str:
        """
        根据文件类型提取文本内容
        """
        try:
            if file_type.lower() == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_from_docx(file_path)
            else:
                # 文本文件直接读取
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
                    
        except Exception as e:
            logger.error(f"文本提取失败: {e}")
            raise ValueError(f"不支持的文件格式或文件损坏: {file_type}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """
        从PDF文件提取文本
        """
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n"
        except Exception as e:
            logger.error(f"PDF解析错误: {e}")
            raise ValueError("PDF文件解析失败")
        
        return text.strip()

    def _extract_from_docx(self, file_path: str) -> str:
        """
        从Word文档提取文本（增强版，支持复杂格式）
        """
        try:
            doc = docx.Document(file_path)
            parts = []
            text_count = 0

            # 1. 段落（主要文本内容）
            for paragraph in doc.paragraphs:
                text = (paragraph.text or '').strip()
                if text:
                    parts.append(text)
                    text_count += len(text)

            # 2. 表格（常见简历内容在表格中）
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = []
                        for p in cell.paragraphs:
                            text = (p.text or '').strip()
                            if text:
                                cell_text.append(text)
                        if cell_text:
                            row_texts.append(' '.join(cell_text))
                    if row_texts:
                        parts.append(' | '.join(row_texts))  # 用 | 分隔表格列
                        text_count += sum(len(t) for t in row_texts)

            # 3. 跳过页眉/页脚（通常包含页码、公司信息等无关内容，不需要提取）

            # 4. 尝试提取所有可能的文本内容（包括嵌套结构）
            try:
                # 递归提取所有文本节点
                def extract_all_text(element):
                    texts = []
                    if hasattr(element, 'text') and element.text:
                        text = element.text.strip()
                        if text:
                            texts.append(text)
                    if hasattr(element, 'iter'):
                        for child in element.iter():
                            if hasattr(child, 'text') and child.text:
                                text = child.text.strip()
                                if text and text not in parts:
                                    texts.append(text)
                    return texts
                
                # 从文档主体提取所有文本
                if hasattr(doc, 'element') and hasattr(doc.element, 'body'):
                    additional_texts = extract_all_text(doc.element.body)
                    if additional_texts:
                        parts.extend(additional_texts)
                        text_count += sum(len(t) for t in additional_texts)
            except Exception as xml_err:
                logger.debug(f"额外文本提取失败（非关键）: {xml_err}")

            # 5. 如果仍然没有提取到足够文本，尝试使用python-docx的完整文本提取
            if text_count < 10:
                try:
                    # 使用docx的完整文本提取方法
                    full_text = []
                    for element in doc.element.body:
                        if hasattr(element, 'iter'):
                            for para in element.iter():
                                if hasattr(para, 'text'):
                                    text = (para.text or '').strip()
                                    if text and text not in parts:
                                        full_text.append(text)
                    if full_text:
                        parts.extend(full_text)
                        text_count += sum(len(t) for t in full_text)
                except Exception as full_err:
                    logger.debug(f"完整文本提取失败: {full_err}")

            extracted_text = "\n".join(parts)
            logger.info(f"[Word提取] 提取到 {len(parts)} 个文本片段，总长度: {len(extracted_text)} 字符")
            
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning(f"[Word提取] 提取的文本过短: {len(extracted_text)} 字符，可能文档格式复杂或为空")
                # 返回提取到的所有文本，即使很短，让AI尝试解析
                return extracted_text.strip() if extracted_text.strip() else "简历文档"
            
            return extracted_text.strip()
            
        except Exception as e:
            logger.error(f"Word文档解析错误: {e}", exc_info=True)
            raise ValueError(f"Word文档解析失败: {str(e)}")

    def _normalize_enhanced_data(self, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        将增强格式转换为兼容格式，同时保留增强信息
        1. 将responsibilities和achievements从对象格式转换为数组格式（兼容旧格式）
        2. 保留增强信息在_enhanced字段中
        3. 处理skills和projects的增强格式
        """
        if not isinstance(structured_data, dict):
            return structured_data
        
        normalized = structured_data.copy()
        enhanced_info = {}  # 存储增强信息
        
        # 处理工作经历
        work_experiences = normalized.get('work_experiences', [])
        if work_experiences:
            normalized_work_exps = []
            for idx, exp in enumerate(work_experiences):
                normalized_exp = exp.copy()
                
                # 处理responsibilities：如果是对象格式，提取raw作为默认值，保留完整对象在_enhanced中
                if isinstance(exp.get('responsibilities'), dict):
                    resp_obj = exp['responsibilities']
                    normalized_exp['responsibilities'] = resp_obj.get('raw', [])
                    if idx not in enhanced_info:
                        enhanced_info[idx] = {}
                    enhanced_info[idx]['responsibilities'] = resp_obj
                elif isinstance(exp.get('responsibilities'), list):
                    # 已经是数组格式，保持不变
                    pass
                
                # 处理achievements：如果是对象格式，提取raw作为默认值
                if isinstance(exp.get('achievements'), dict):
                    ach_obj = exp['achievements']
                    normalized_exp['achievements'] = ach_obj.get('raw', [])
                    if idx not in enhanced_info:
                        enhanced_info[idx] = {}
                    enhanced_info[idx]['achievements'] = ach_obj
                elif isinstance(exp.get('achievements'), list):
                    # 已经是数组格式，保持不变
                    pass
                
                # 处理skills_used：如果是对象格式，合并explicit和implicit
                if isinstance(exp.get('skills_used'), dict):
                    skills_obj = exp['skills_used']
                    normalized_exp['skills_used'] = (
                        skills_obj.get('explicit', []) + 
                        skills_obj.get('implicit', [])
                    )
                    if idx not in enhanced_info:
                        enhanced_info[idx] = {}
                    enhanced_info[idx]['skills_used'] = skills_obj
                elif isinstance(exp.get('skills_used'), list):
                    # 已经是数组格式，保持不变
                    pass
                
                # 保留implicit_info、related_projects、_paragraphs等增强字段
                # 这些字段不影响兼容性，直接保留
                
                normalized_work_exps.append(normalized_exp)
            
            normalized['work_experiences'] = normalized_work_exps
            if enhanced_info:
                normalized['_work_experiences_enhanced'] = enhanced_info
        
        # 处理教育背景：修正学历和学位的逻辑错误
        education = normalized.get('education', [])
        if education:
            for edu in education:
                if not isinstance(edu, dict):
                    continue
                degree = edu.get('degree', '') or ''
                education_level = edu.get('education_level', '') or ''
                
                # 如果学位是"硕士"或"博士"，但学历是"本科"，修正为"研究生"或空
                if degree in ['硕士', '博士']:
                    if education_level == '本科':
                        edu['education_level'] = '研究生'
                        logger.info(f"[数据规范化] 修正教育背景：degree={degree}，将 education_level 从'本科'修正为'研究生'")
                    elif education_level in ['硕士', '博士']:
                        # 如果学历错误地填入了"硕士"或"博士"，清空或设为"研究生"
                        edu['education_level'] = '研究生'
                        logger.info(f"[数据规范化] 修正教育背景：degree={degree}，将 education_level 从'{education_level}'修正为'研究生'")
        
        # 处理技能
        skills = normalized.get('skills', {})
        if isinstance(skills, dict):
            normalized_skills = {}
            skills_enhanced = {}
            
            # 处理technical技能
            if isinstance(skills.get('technical'), dict):
                tech_obj = skills['technical']
                normalized_skills['technical'] = (
                    tech_obj.get('explicit', []) + 
                    tech_obj.get('inferred', [])
                )
                skills_enhanced['technical'] = tech_obj
            elif isinstance(skills.get('technical'), list):
                normalized_skills['technical'] = skills['technical']
            
            # soft和languages保持不变
            if 'soft' in skills:
                normalized_skills['soft'] = skills['soft']
            if 'languages' in skills:
                normalized_skills['languages'] = skills['languages']
            
            normalized['skills'] = normalized_skills
            if skills_enhanced:
                normalized['_skills_enhanced'] = skills_enhanced
        
        # 处理项目
        projects = normalized.get('projects', [])
        if projects:
            normalized_projects = []
            projects_enhanced = {}
            
            for idx, proj in enumerate(projects):
                normalized_proj = proj.copy()
                
                # 处理description：统一为字符串格式
                desc_value = proj.get('description')
                if isinstance(desc_value, dict):
                    # 如果是对象格式，提取raw作为默认值
                    desc_obj = desc_value
                    normalized_proj['description'] = desc_obj.get('raw', '')
                    if idx not in projects_enhanced:
                        projects_enhanced[idx] = {}
                    projects_enhanced[idx]['description'] = desc_obj
                elif isinstance(desc_value, list):
                    # 如果是数组，合并成字符串（清理特殊字符后合并）
                    normalized_proj['description'] = ' '.join([str(d).strip() for d in desc_value if d and str(d).strip()])
                    logger.info(f"[数据规范化] 项目{idx}的description从数组转换为字符串: {len(desc_value)}项")
                elif isinstance(desc_value, str):
                    # 已经是字符串格式，保持不变
                    normalized_proj['description'] = desc_value
                elif desc_value is not None:
                    # 其他类型，转换为字符串
                    normalized_proj['description'] = str(desc_value)
                else:
                    normalized_proj['description'] = ''
                
                # 处理achievements：如果是对象格式，提取raw作为默认值
                if isinstance(proj.get('achievements'), dict):
                    ach_obj = proj['achievements']
                    normalized_proj['achievements'] = ach_obj.get('raw', '')
                    if idx not in projects_enhanced:
                        projects_enhanced[idx] = {}
                    projects_enhanced[idx]['achievements'] = ach_obj
                elif isinstance(proj.get('achievements'), str):
                    # 已经是字符串格式，保持不变
                    pass
                
                # 保留related_work等增强字段
                normalized_projects.append(normalized_proj)
            
            normalized['projects'] = normalized_projects
            if projects_enhanced:
                normalized['_projects_enhanced'] = projects_enhanced
        
        # 保留_metadata等元数据
        if '_metadata' in normalized:
            # _metadata字段不影响兼容性，直接保留
            pass
        
        return normalized
    
    def _apply_dynamic_field_cleanup(self, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        根据实际内容清理解析结果：移除空字段，保留真实存在的信息
        """
        if not isinstance(structured_data, dict):
            return structured_data

        cleaned = structured_data.copy()

        # 基本信息：仅保留有值的字段
        if isinstance(cleaned.get("basic_info"), dict):
            cleaned["basic_info"] = self._remove_empty_values(cleaned["basic_info"])

        # 技能字段：移除空的技能分类
        if isinstance(cleaned.get("skills"), dict):
            cleaned["skills"] = self._remove_empty_values(cleaned["skills"])

        # 工作经历、教育、项目：逐条清理空字段，并移除空记录
        for section in ["work_experiences", "education", "projects"]:
            if isinstance(cleaned.get(section), list):
                cleaned_section = []
                for item in cleaned[section]:
                    if not isinstance(item, dict):
                        continue
                    cleaned_item = self._remove_empty_values(item)
                    if not self._is_empty_value(cleaned_item):
                        cleaned_section.append(cleaned_item)
                cleaned[section] = cleaned_section

        # 其他简单字段（如 professional_summary）去除首尾空白
        if isinstance(cleaned.get("professional_summary"), str):
            cleaned["professional_summary"] = cleaned["professional_summary"].strip()

        return cleaned

    def _remove_empty_values(self, value: Any) -> Any:
        """
        递归移除空值（None、空字符串、空列表、空字典）
        """
        if isinstance(value, dict):
            result = {}
            for key, val in value.items():
                cleaned_val = self._remove_empty_values(val)
                if self._is_empty_value(cleaned_val):
                    continue
                result[key] = cleaned_val
            return result
        if isinstance(value, list):
            result_list = []
            for item in value:
                cleaned_item = self._remove_empty_values(item)
                if self._is_empty_value(cleaned_item):
                    continue
                result_list.append(cleaned_item)
            return result_list
        if isinstance(value, str):
            return value.strip()
        return value

    def _is_empty_value(self, value: Any) -> bool:
        """
        判断值是否为空，用于清理动态字段
        """
        if value is None:
            return True
        if isinstance(value, bool):
            return False
        if isinstance(value, str):
            return value == ""
        if isinstance(value, (list, tuple, set)):
            return len(value) == 0
        if isinstance(value, dict):
            return len(value) == 0
        return False

    async def validate_parsed_data(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        验证解析数据的完整性和质量，主动识别问题
        """
        validation_result = {
            "is_valid": True,
            "missing_fields": [],
            "data_quality_issues": [],
            "suggestions": [],
            "warnings": []
        }
        
        # 检查基本信息
        basic_info = parsed_data.get('basic_info', {})
        if not basic_info.get('name'):
            validation_result["missing_fields"].append("姓名")
            validation_result["is_valid"] = False
        if not basic_info.get('email'):
            validation_result["missing_fields"].append("邮箱")
            validation_result["is_valid"] = False
        if not basic_info.get('phone'):
            validation_result["warnings"].append("缺少电话，建议补充")
        
        # 检查工作经历
        work_experiences = parsed_data.get('work_experiences', [])
        if not work_experiences:
            validation_result["suggestions"].append("建议补充工作经历")
        else:
            # 检查每段工作经历的完整性
            for idx, exp in enumerate(work_experiences, 1):
                if not exp.get('company'):
                    validation_result["data_quality_issues"].append(f"第{idx}段工作经历缺少公司名称")
                if not exp.get('position'):
                    validation_result["data_quality_issues"].append(f"第{idx}段工作经历缺少职位")
                if not exp.get('start_date'):
                    validation_result["data_quality_issues"].append(f"第{idx}段工作经历缺少开始时间")
                if not exp.get('responsibilities') or len(exp.get('responsibilities', [])) == 0:
                    validation_result["warnings"].append(f"第{idx}段工作经历缺少工作职责描述")
                if not exp.get('achievements') or len(exp.get('achievements', [])) == 0:
                    validation_result["warnings"].append(f"第{idx}段工作经历缺少工作成就，建议补充量化成果")
            
            # 检查时间逻辑
            for i in range(len(work_experiences) - 1):
                current = work_experiences[i]
                next_exp = work_experiences[i + 1]
                current_end = current.get('end_date') or (current.get('start_date') if current.get('is_current') else None)
                next_start = next_exp.get('start_date')
                
                if current_end and next_start and current_end < next_start:
                    # 时间倒序是正常的（最新的在前），但如果差距太大可能是问题
                    pass  # 暂时不标记为错误
                elif current_end and next_start and current_end > next_start:
                    validation_result["data_quality_issues"].append(f"工作经历时间可能重叠：第{i+1}段结束时间({current_end})晚于第{i+2}段开始时间({next_start})")
        
        # 检查教育背景
        education = parsed_data.get('education', [])
        if not education:
            validation_result["suggestions"].append("建议补充教育背景")
        else:
            for idx, edu in enumerate(education, 1):
                if not edu.get('school'):
                    validation_result["data_quality_issues"].append(f"第{idx}个教育背景缺少学校名称")
                if not edu.get('major'):
                    validation_result["warnings"].append(f"第{idx}个教育背景缺少专业信息")
        
        # 检查技能
        skills = parsed_data.get('skills', {})
        technical_skills = skills.get('technical', [])
        if not technical_skills or len(technical_skills) == 0:
            validation_result["warnings"].append("缺少技术技能，建议补充")
        
        # 检查项目经验
        projects = parsed_data.get('projects', [])
        if not projects:
            validation_result["suggestions"].append("建议补充项目经验，可以提升简历竞争力")
        
        # 生成总结性建议
        if validation_result["missing_fields"]:
            validation_result["suggestions"].insert(0, f"缺少关键信息：{', '.join(validation_result['missing_fields'])}，请补充")
        if validation_result["data_quality_issues"]:
            validation_result["suggestions"].append(f"发现{len(validation_result['data_quality_issues'])}个数据质量问题，请检查并修正")
            
        return validation_result
    
    def _preprocess_text(self, text: str) -> str:
        """
        预处理文本：清理和优化，减少AI处理负担
        """
        if not text:
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # 移除行首行尾空白
            line = line.strip()
            
            # 跳过空行
            if not line:
                continue
            
            # 跳过明显的页眉页脚（页码、日期等）
            # 但要谨慎，避免误删重要内容
            if self._is_header_footer(line):
                # 记录被跳过的内容，便于调试
                logger.debug(f"[文本预处理] 跳过可能的页眉页脚: {line[:50]}...")
                continue
            
            # 移除过多的连续空格
            line = ' '.join(line.split())
            
            # 移除过短的无意义行（可能是格式字符）
            if len(line) < 2:
                continue
            
            cleaned_lines.append(line)
        
        # 合并为文本
        cleaned_text = '\n'.join(cleaned_lines)
        original_length = len(cleaned_text)
        
        # 提高截断限制，并智能截断（在段落边界截断）
        max_length = 30000  # 提高到30000字符，确保大部分简历完整处理
        
        if len(cleaned_text) > max_length:
            logger.warning(f"[文本预处理] 文本过长({len(cleaned_text)}字符)，将智能截断至{max_length}字符")
            # 智能截断：在段落边界截断，而不是简单截断
            cleaned_text = self._smart_truncate(cleaned_text, max_length)
            logger.warning(f"[文本预处理] 截断后长度: {len(cleaned_text)}字符，丢失: {original_length - len(cleaned_text)}字符")
        else:
            logger.info(f"[文本预处理] 文本长度: {original_length}字符，无需截断")
        
        return cleaned_text
    
    def _smart_truncate(self, text: str, max_length: int) -> str:
        """
        智能截断：在段落边界截断，保留关键信息
        优先保留：工作经历、项目经历、教育背景等关键部分
        """
        if len(text) <= max_length:
            return text
        
        # 尝试在段落边界截断（以换行符为界）
        # 先尝试保留前max_length-100字符，然后找到最后一个完整的段落
        target_length = max_length - 100  # 留出100字符的缓冲
        truncated = text[:target_length]
        
        # 查找最后一个换行符（段落边界）
        last_newline = truncated.rfind('\n')
        if last_newline > target_length * 0.8:  # 如果最后一个换行符在80%位置之后，使用它
            truncated = truncated[:last_newline]
        else:
            # 如果找不到合适的换行符，尝试查找句号、分号等句子边界
            sentence_endings = ['。', '；', '. ', '; ', '\n\n']
            for ending in sentence_endings:
                last_ending = truncated.rfind(ending)
                if last_ending > target_length * 0.8:
                    truncated = truncated[:last_ending + len(ending)]
                    break
        
        # 添加截断标记
        truncated += f"\n\n[注意：文本已截断，原始长度{len(text)}字符，当前保留{len(truncated)}字符]"
        
        return truncated
    
    def _is_header_footer(self, line: str) -> bool:
        """
        判断是否是页眉页脚
        注意：必须非常谨慎，避免误删工作经历等重要内容
        """
        import re
        line_lower = line.lower()
        line_stripped = line.strip()
        
        # 白名单：如果包含工作经历相关关键词，绝对不是页眉页脚
        work_experience_keywords = [
            '公司', '职位', '岗位', '部门', '负责', '职责', '工作', '项目',
            '业绩', '成就', '成果', '管理', '团队', '下属', '汇报',
            'company', 'position', 'responsibility', 'work', 'project'
        ]
        if any(keyword in line for keyword in work_experience_keywords):
            return False
        
        # 白名单：如果包含日期范围（如 2019/4—2022/6），绝对不是页眉页脚
        # 日期范围通常包含两个日期，用破折号、横线等连接
        date_range_patterns = [
            r'\d{4}[-/\.]\d{1,2}[-—–]\d{4}[-/\.]\d{1,2}',  # 2019/4—2022/6, 2019.4-2022.6
            r'\d{4}年\d{1,2}月[-—–]\d{4}年\d{1,2}月',      # 2019年4月—2022年6月
        ]
        for pattern in date_range_patterns:
            if re.search(pattern, line):
                return False
        
        # 页码模式（更严格）
        if any(pattern in line_lower for pattern in ['第', '页', 'page']):
            # 排除包含"第X年"、"第X月"这样的日期
            if '年' not in line and '月' not in line:
                # 进一步检查：如果包含数字+页，且行长度很短，才认为是页码
                if re.match(r'^第?\d+页?$', line_stripped) or re.match(r'^page\s*\d+$', line_lower):
                    return True
        
        # 日期模式（非常严格：只匹配"单独一行只有日期"的情况）
        # 必须是整行就是一个日期，且不包含其他文字
        single_date_patterns = [
            r'^\d{4}[-/]\d{1,2}[-/]\d{1,2}$',      # 2024-01-01（整行只有这个）
            r'^\d{4}年\d{1,2}月\d{1,2}日$',        # 2024年1月1日（整行只有这个）
            r'^\d{4}[-/]\d{1,2}$',                  # 2024-01（整行只有这个）
            r'^\d{4}年\d{1,2}月$',                  # 2024年1月（整行只有这个）
        ]
        for pattern in single_date_patterns:
            if re.match(pattern, line_stripped):
                # 再次确认：整行只有日期，没有其他内容
                if len(line_stripped) <= 15:  # 日期格式最长约15字符
                    return True
        
        return False

# 全局解析器实例
resume_parser = ResumeParser()