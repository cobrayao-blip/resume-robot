from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
import logging
import os
import re
from datetime import datetime
from ..core.config import settings

logger = logging.getLogger(__name__)

class WordExporter:
    def __init__(self):
        self.doc = Document()
        
    def export_resume(self, resume_data: Dict[str, Any], template_config: Dict[str, Any] = None) -> str:
        """
        导出简历为Word文档
        """
        try:
            # 每次导出都重新创建文档，防止上一次内容残留导致重复
            self.doc = Document()
            # 设置文档基本样式
            self._setup_document_styles()
            
            # 文档标题固定为"人才推荐报告"
            title_text = '人才推荐报告'
            # 使用普通段落而不是heading，避免默认的蓝色下划线
            title_para = self.doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(title_text)
            title_run.font.bold = True
            title_run.font.size = Pt(16)  # 小二字体（16pt）
            title_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 设置中文字体为黑体
            title_run.font.name = '黑体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置段前间距（如12磅）
            title_para.paragraph_format.space_before = Pt(12)
            # 设置段后间距（如10磅）
            title_para.paragraph_format.space_after = Pt(10)
            # self.doc.add_paragraph()
            
            # 严格按 template_sections 顺序渲染，每个组件使用自己的 fields 和 data
            raw_sections = resume_data.get('template_sections') if isinstance(resume_data, dict) else None
            sections = raw_sections if isinstance(raw_sections, list) else []
            if sections:
                section_index = 0
                for sec in sections:
                    try:
                        if not isinstance(sec, dict):
                            logger.warning("Skip invalid section (not dict): %r", sec)
                            continue
                        sec_type = sec.get('type')
                        sec_title = sec.get('title')
                        cfg = sec.get('config') or {}
                        fields = sec.get('fields') or []
                        # 仅接受列表或字典的字段数组，其他类型丢弃
                        if not isinstance(fields, list):
                            fields = []
                        sec_data = sec.get('data')

                        # 检查该 section 是否需要渲染（项目经历如果没有内容就不渲染）
                        should_render = True
                        if sec_type == 'projects':
                            # 项目经历：检查是否有数据
                            rows = []
                            if isinstance(sec_data, dict) and isinstance(sec_data.get('rows'), list):
                                rows = sec_data.get('rows')
                            elif isinstance(sec_data, list):
                                rows = sec_data
                            # 如果没有数据，也不渲染（整个组件不显示）
                            if not rows or len(rows) == 0:
                                should_render = False
                        
                        # 如果不需要渲染，跳过该 section（不增加 section_index）
                        if not should_render:
                            continue

                        # 章节级数字编号：1、2、3 ...（只有需要渲染的 section 才增加编号）
                        section_index += 1
                        self._add_numbered_section_title(section_index, sec_title or sec_type or '章节')

                        if sec_type == 'basic_info':
                            self._add_basic_info_by_fields(sec_title or '基本信息', fields, sec_data, None, add_title=False)
                        elif sec_type == 'education':
                            self._add_education_by_fields(sec_title or '教育背景', cfg, fields, sec_data, None, add_title=False)
                        elif sec_type == 'work_experience':
                            # 支持 displayMode 配置：summary（仅表格）、detailed（仅详细卡片）、both（表格+详细卡片）
                            self._add_work_experience_by_fields(sec_title or '工作经历', cfg, fields, sec_data, None, add_title=False)
                        elif sec_type == 'recommended_jobs':
                            self._add_recommended_jobs({
                                'items': (sec_data.get('rows') if isinstance(sec_data, dict) else (sec_data if isinstance(sec_data, list) else None)) or [],
                                'view': cfg.get('view') or 'list',
                                'fields': fields,
                                'title': sec_title or '推荐岗位'
                            }, add_title=False)
                        elif sec_type == 'evaluation':
                            self._add_evaluation_by_fields(sec_title or '评价', fields, sec_data, None, add_title=False)
                        elif sec_type == 'salary':
                            self._add_salary_by_fields(sec_title or '薪资', fields, sec_data, None, add_title=False)
                        elif sec_type == 'skills':
                            self._add_skills_by_fields(sec_title or '技能专长', fields, sec_data, None, add_title=False)
                        elif sec_type == 'projects':
                            # 传入原始 projects 作为后备数据，以便在字段数据中补充 start_date/end_date 等时间信息
                            self._add_projects_by_fields(
                                sec_title or '项目经历',
                                fields,
                                sec_data,
                                resume_data.get('projects', []),
                                add_title=False
                            )
                        else:
                            # 未知组件：按字段占位渲染
                            self._add_generic_section(sec_title or sec_type or '未知组件', fields, sec_data, add_title=False)
                    except Exception as sec_err:
                        logger.error("Render section failed: type=%r title=%r error=%s", sec.get('type') if isinstance(sec, dict) else None, sec.get('title') if isinstance(sec, dict) else None, sec_err)
                        # 出错不断导，继续其它 section
                        continue
            else:
                # 默认顺序
                self._add_basic_info(resume_data.get('basic_info', {}))
                self._add_education(resume_data.get('education', []))
                self._add_work_experience(resume_data.get('work_experiences', []))
                self._add_recommended_jobs(resume_data.get('recommended_jobs'))
                self._add_evaluation(resume_data.get('evaluation'))
                self._add_salary(resume_data.get('salary'))
                self._add_skills(resume_data.get('skills', {}))
                self._add_projects(resume_data.get('projects', []))
            
            # 生成文件名
            filename = self._generate_filename(resume_data)
            
            return filename
            
        except Exception as e:
            logger.error(f"Word导出失败: {e}")
            raise Exception(f"文档生成失败: {str(e)}")
    
    def _format_date(self, date_str: str) -> str:
        """
        格式化日期：将 YYYY-MM 格式转换为 YYYY.MM 格式
        """
        if not date_str:
            return date_str
        # 将 YYYY-MM 格式转换为 YYYY.MM
        return date_str.replace('-', '.')
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本中的特殊字符，避免段落零碎
        移除：↓、↩、以及其他可能导致格式问题的特殊字符
        处理换行符：将句子中间的单个换行符替换为空格，只保留段落分隔
        注意：此方法用于Word导出阶段的最终格式化，数据填充阶段应该已经清理过
        """
        if not text or not isinstance(text, str):
            return text if text else ''
        
        # 移除特殊字符：↓、↩、以及其他格式标记字符
        # ↓ (U+2193), ↩ (U+21A9), ↲ (U+21B2), ↳ (U+21B3)
        # 以及其他可能的格式标记
        special_chars = [
            '\u2193',  # ↓
            '\u21A9',  # ↩
            '\u21B2',  # ↲
            '\u21B3',  # ↳
            '\u2191',  # ↑
            '\u2192',  # →
            '\u2190',  # ←
            '\u21E8',  # ⇨
            '\u21E6',  # ⇦
            '\u21E7',  # ⇧
            '\u21E9',  # ⇩
        ]
        
        cleaned = text
        for char in special_chars:
            cleaned = cleaned.replace(char, ' ')
        
        # 处理换行符：将句子中间的单个换行符替换为空格
        # 1. 先处理多个连续换行（段落分隔），保留为双换行
        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)
        # 2. 将单个换行符（前后不是换行符的）替换为空格
        # 但保留项目符号后的换行（如 "● " 或 "• " 后的换行）
        # 先标记项目符号后的换行
        cleaned = re.sub(r'([●•·])\s*\n\s*', r'\1 ', cleaned)
        # 将剩余的单个换行符替换为空格
        cleaned = re.sub(r'(?<!\n)\n(?!\n)', ' ', cleaned)
        # 3. 移除多余的连续空格
        cleaned = re.sub(r' +', ' ', cleaned)
        # 4. 移除行尾的空格
        cleaned = re.sub(r' +$', '', cleaned, flags=re.MULTILINE)
        # 5. 清理段落分隔周围的空格
        cleaned = re.sub(r' \n\n ', '\n\n', cleaned)
        cleaned = re.sub(r'\n\n +', '\n\n', cleaned)
        cleaned = re.sub(r' +\n\n', '\n\n', cleaned)
        
        return cleaned.strip()
    
    def _setup_document_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(10.5)
        
        # 设置中文字体
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def _add_numbered_section_title(self, index: int, title: str):
        p = self.doc.add_paragraph()
        run_num = p.add_run(f"{index}.")
        run_space = p.add_run(" ")  # 数字和标题之间的空格
        run_title = p.add_run(title)
        # 数字编号和标题使用相同的字体、字号和加粗
        run_num.font.bold = True
        run_num.font.size = Pt(12)
        run_title.font.bold = True
        run_title.font.size = Pt(12)
        # 空行
        # self.doc.add_paragraph()
    
    def _add_basic_info(self, basic_info: Dict[str, Any]):
        """添加基本信息（与预览一致：作为独立章节，字段名前加粗，值正常，使用项目符号）"""
        if not basic_info:
            return
        self._add_section_title('基本信息')

        def add_item(label: str, value: str):
            if not value:
                return
            p = self.doc.add_paragraph()
            p.style = 'List Bullet'
            run_label = p.add_run(f"{label}：")
            run_label.font.bold = True
            p.add_run(value)

        add_item('姓名', basic_info.get('name', ''))
        add_item('电话', basic_info.get('phone', ''))
        add_item('邮箱', basic_info.get('email', ''))
        add_item('所在地', basic_info.get('location', ''))
        add_item('出生年月', basic_info.get('birth_date', '') or basic_info.get('birthday', ''))
        # 链接使用可点击样式
        website = basic_info.get('website') or basic_info.get('website_url')
        if website:
            p = self.doc.add_paragraph(); p.style = 'List Bullet'
            r1 = p.add_run('网站：'); r1.font.bold = True
            p.add_run(website)
        github = basic_info.get('github')
        if github:
            p = self.doc.add_paragraph(); p.style = 'List Bullet'
            r1 = p.add_run('GitHub：'); r1.font.bold = True
            p.add_run(github)
        linkedin = basic_info.get('linkedin') or basic_info.get('linkedin_url')
        if linkedin:
            p = self.doc.add_paragraph(); p.style = 'List Bullet'
            r1 = p.add_run('LinkedIn：'); r1.font.bold = True
            p.add_run(linkedin)
        self.doc.add_paragraph()
    
    def _add_professional_summary(self, summary: str):
        """添加职业摘要"""
        if summary:
            self._add_section_title('职业摘要')
            cleaned_summary = self._clean_text(summary)
            summary_paragraph = self.doc.add_paragraph()
            summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            summary_run = summary_paragraph.add_run(cleaned_summary)
            summary_run.font.size = Pt(10.5)
            self.doc.add_paragraph()  # 空行
    
    def _add_work_experience(self, work_experiences: List[Dict[str, Any]]):
        """添加工作经历"""
        if not work_experiences:
            return
            
        self._add_section_title('工作经历')
        
        # 添加职业履历表格
        self._add_work_experience_table(work_experiences)
        self.doc.add_paragraph()  # 空行
        
        # 添加详细工作经历
        for i, exp in enumerate(work_experiences, 1):
            self._add_detailed_experience(exp, i)
    
    def _add_work_experience_table(self, experiences: List[Dict[str, Any]]):
        """添加工作经历表格"""
        if not experiences:
            return
            
        # 创建表格（起止时间、公司名称、职位）
        table = self.doc.add_table(rows=len(experiences) + 1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头（与预览一致）
        headers = ['起止时间', '公司名称', '职位']
        for col, header in enumerate(headers):
            cell = table.cell(0, col)
            cell.text = header
            # 设置表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 填充数据
        for row, exp in enumerate(experiences, 1):
            # 起止时间
            start_date = self._format_date(exp.get('start_date', ''))
            end_date = self._format_date(exp.get('end_date', ''))
            is_current = exp.get('is_current', False)
            period = f"{start_date} - {'至今' if is_current else end_date}"
            table.cell(row, 0).text = period
            # 公司名称
            company = self._clean_text(exp.get('company', ''))
            table.cell(row, 1).text = company
            # 职位
            position = self._clean_text(exp.get('position', ''))
            table.cell(row, 2).text = position
    
    def _add_detailed_experience(self, experience: Dict[str, Any], index: int):
        """添加详细工作经历 - 与前端预览的工作详情排版保持一致"""
        company = experience.get('company', '')
        position = experience.get('position', '')
        start_date = self._format_date(experience.get('start_date', ''))
        end_date = self._format_date(experience.get('end_date', ''))
        is_current = experience.get('is_current', False)
        location = experience.get('location', '')
        report_to = experience.get('report_to', '')
        team_size = experience.get('team_size', '')

        # 1) 标题行：起止时间, 公司名称（加粗+下划线）
        period_text = f"{start_date} - {'至今' if is_current else end_date}".strip()
        title_text = ", ".join([t for t in [period_text, company] if t])
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run(title_text if title_text else company)
        title_run.font.bold = True
        title_run.font.underline = True
        title_run.font.size = Pt(11)

        # 2) 带 bullet 的明细（字段名加粗，值缩进）
        # 职位、汇报对象、下属团队、工作地点
        bullet_items = [
            ("职位", position),
            ("汇报对象", report_to),
            ("下属团队", str(team_size) if team_size else ''),
            ("工作地点", location),
        ]
        for label, value in bullet_items:
            if not value:
                continue
            p = self.doc.add_paragraph()
            p.style = 'List Bullet'
            run_label = p.add_run(f"{label}: ")
            run_label.font.bold = True
            cleaned_value = self._clean_text(str(value)) if value else ''
            p.add_run(cleaned_value)

        # 3) 长文本：工作职责、工作业绩、离职原因（每个字段从1开始手动编号，字段内连续编号）
        # 与前端模板组件保持一致：使用HTML的<ol>标签，每个<ol>的编号从1开始
        numbered_blocks = [
            ("工作职责", experience.get('responsibilities', [])),
            ("工作业绩", experience.get('achievements', [])),
            ("离职原因", [experience.get('reason_for_leaving', '')] if experience.get('reason_for_leaving') else []),
        ]
        for label, items in numbered_blocks:
            if not items:
                continue
            # 添加标签（粗体），使用 List Bullet 样式，与职位等字段保持一致
            p = self.doc.add_paragraph()
            p.style = 'List Bullet'
            run = p.add_run(f"{label}:")
            run.font.bold = True
            # 手动添加编号，每个字段从1开始（与前端<ol>标签行为一致）
            for idx, item in enumerate(items, 1):
                if not item or not str(item).strip():
                    continue
                # 检查item内容是否已经包含编号，如果包含则移除（避免重复编号）
                item_str = str(item).strip()
                # 移除开头的数字编号格式（如"1. "、"1、"、"1．"、"1)"等）
                # 支持阿拉伯数字+常见分隔符（半角/全角点、中文顿号、右括号）
                item_str = re.sub(r'^\s*\d+\s*[\.、．\)]\s*', '', item_str)
                # 清理特殊字符
                item_str = self._clean_text(item_str)
                
                li = self.doc.add_paragraph()
                li.style = 'Normal'
                # 设置悬挂缩进：第一行文字可以在编号后有适当空间，第二行及后续行与第一行文字对齐
                # left_indent: 控制第二行及后续行的起始位置（与第一行文字对齐）
                # first_line_indent: 负值表示悬挂缩进，让第一行向左突出，容纳编号和间距
                # 编号"1. "大约占0.3英寸，编号后间距约0.1英寸，所以第一行文字起始位置约0.4英寸
                # 第二行及后续行应该与第一行文字对齐，所以left_indent设为0.4英寸
                # first_line_indent设为-0.15英寸，让第一行向左突出，容纳编号
                li.paragraph_format.left_indent = Inches(0.4)  # 第二行及后续行的起始位置（与第一行文字对齐）
                li.paragraph_format.first_line_indent = Inches(-0.15)  # 悬挂缩进，让第一行向左突出以容纳编号
                # 设置段后间距6磅
                li.paragraph_format.space_after = Pt(6)
                # 手动添加编号文本（"1. "、"2. "等）
                li.add_run(f"{idx}. {item_str}")

        # self.doc.add_paragraph()  # 空行
    
    def _add_education(self, education: List[Dict[str, Any]]):
        """添加教育背景为表格（与前端预览一致）"""
        if not education:
            return
        
        self._add_section_title('教育背景（从最高学历开始）')
        
        headers = ['起止时间', '学校', '专业', '学历', '学位', '备注']
        table = self.doc.add_table(rows=len(education) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
        
        # 数据
        for row, edu in enumerate(education, start=1):
            period = edu.get('period')
            if period:
                # 如果period已存在，也需要格式化（可能包含 YYYY-MM 格式）
                period = self._format_date(period) if '-' in period else period
            else:
                period = self._format_date(edu.get('start_date', ''))
            if 'end_date' in edu:
                start = self._format_date(edu.get('start_date', ''))
                end = self._format_date(edu.get('end_date', ''))
                period = f"{start} - {end}"
            elif 'graduation_date' in edu and not period:
                period = self._format_date(edu.get('graduation_date', ''))
            table.cell(row, 0).text = period
            table.cell(row, 1).text = edu.get('school', '')
            table.cell(row, 2).text = edu.get('major', '')
            table.cell(row, 3).text = edu.get('education_level', '') or edu.get('degree_level', '')
            table.cell(row, 4).text = edu.get('degree', '')
            table.cell(row, 5).text = edu.get('remark', '')
        
        self.doc.add_paragraph()  # 空行
    
    def _add_skills(self, skills: Dict[str, Any]):
        """添加技能专长"""
        if not skills:
            return
            
        self._add_section_title('技能专长')
        
        # 技术技能
        technical_skills = skills.get('technical', [])
        if technical_skills:
            tech_paragraph = self.doc.add_paragraph()
            tech_run = tech_paragraph.add_run("技术技能: ")
            tech_run.font.bold = True
            tech_run = tech_paragraph.add_run(", ".join(technical_skills))
        
        # 软技能
        soft_skills = skills.get('soft', [])
        if soft_skills:
            soft_paragraph = self.doc.add_paragraph()
            soft_run = soft_paragraph.add_run("软技能: ")
            soft_run.font.bold = True
            soft_run = soft_paragraph.add_run(", ".join(soft_skills))
        
        # 语言能力
        languages = skills.get('languages', [])
        if languages:
            lang_paragraph = self.doc.add_paragraph()
            lang_run = lang_paragraph.add_run("语言能力: ")
            lang_run.font.bold = True
            lang_run = lang_paragraph.add_run(", ".join(languages))
        
        self.doc.add_paragraph()  # 空行
    
    def _add_projects(self, projects: List[Dict[str, Any]]):
        """添加项目经历"""
        if not projects:
            return
            
        self._add_section_title('项目经历')
        
        for idx, project in enumerate(projects):
            # 第二个及后续项目：先添加空段落（换行），确保段前间距生效
            if idx > 0:
                self.doc.add_paragraph()
            
            project_paragraph = self.doc.add_paragraph()
            
            # 项目名称
            name = project.get('name', '')
            project_run = project_paragraph.add_run(name)
            project_run.font.bold = True
            
            # 担任角色
            role = project.get('role', '')
            if role:
                role_run = project_paragraph.add_run(f" ({role})")
                role_run.font.color.rgb = RGBColor(100, 100, 100)

            # 项目起止时间（使用基础数据中的 start_date/end_date，导出为 YYYY.MM）
            start_date = project.get('start_date') or ''
            end_date = project.get('end_date') or ''
            start_formatted = self._format_date(start_date) if start_date else ''
            end_formatted = self._format_date(end_date) if end_date else ''

            period = ''
            if start_formatted and end_formatted:
                period = f"{start_formatted} - {end_formatted}"
            elif start_formatted and not end_formatted:
                period = f"{start_formatted} - 至今"
            elif not start_formatted and end_formatted:
                period = end_formatted

            if period:
                period_run = project_paragraph.add_run(f"  {period}")
                period_run.font.color.rgb = RGBColor(120, 120, 120)
            
            # 设置项目名称段落格式：第一个项目不设置段前间距，后续项目设置段前间距12磅
            # 先清除可能的默认值，再设置新值，确保设置生效
            if idx > 0:
                project_paragraph.paragraph_format.space_before = Pt(12)
            # 设置项目名称段落段后间距10磅
            project_paragraph.paragraph_format.space_after = Pt(10)
                                   
            # 项目描述
            description = project.get('description', '')
            if description:
                # 处理description：如果是数组，合并成字符串；如果是字符串，直接使用
                if isinstance(description, list):
                    # 如果是数组，合并成字符串（用空格连接，避免引入换行）
                    description = ' '.join([str(d).strip() for d in description if d and str(d).strip()]) if description else ''
                elif not isinstance(description, str):
                    description = str(description) if description else ''
                
                if description:
                    cleaned_description = self._clean_text(description)
                    desc_paragraph = self.doc.add_paragraph(cleaned_description)
                    desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    # 设置段后间距6磅
                    desc_paragraph.paragraph_format.space_after = Pt(6)
            
            # 项目业绩：优先使用 achievements，如果没有再使用 outcome
            achievements = project.get('achievements', '')
            outcome = project.get('outcome', '')
            
            # 处理 achievements：如果是数组，合并成字符串
            if isinstance(achievements, list):
                achievements_text = ' '.join([str(a).strip() for a in achievements if a and str(a).strip()])
            elif isinstance(achievements, str):
                achievements_text = achievements
            else:
                achievements_text = ''
            
            # 合并：优先使用 achievements，如果没有再使用 outcome
            final_achievements = achievements_text or outcome
            
            if final_achievements:
                cleaned_achievements = self._clean_text(final_achievements)
                achievements_paragraph = self.doc.add_paragraph()
                achievements_run = achievements_paragraph.add_run("项目业绩: ")
                achievements_run.font.bold = True
                achievements_run = achievements_paragraph.add_run(cleaned_achievements)
                # 设置段后间距6磅
                achievements_paragraph.paragraph_format.space_after = Pt(6)

    def _add_recommended_jobs(self, jobs: Any, add_title: bool = True):
        """添加推荐岗位（列表/表格）"""
        if jobs is None:
            return
        title = '推荐岗位'
        if isinstance(jobs, dict) and jobs.get('title'):
            title = str(jobs.get('title'))
        if add_title:
            self._add_section_title(title)
        # 兼容：list 或 { items: [...], view: 'list'|'table' }
        rows = []
        view = 'table'
        fields = []
        if isinstance(jobs, list):
            rows = jobs
        elif isinstance(jobs, dict):
            cand = jobs.get('items') or jobs.get('rows') or jobs.get('data')
            if isinstance(cand, list):
                rows = cand
            view = jobs.get('view') or 'table'
            fields = jobs.get('fields') or []
        # 当无数据时也保留版式：渲染占位
        if view == 'list':
            # 以列表显示：每个岗位的所有字段都使用相同的项目符号对齐
            items = rows if rows else [{}]
            for job in items:
                # 所有字段使用相同的缩进，确保项目符号对齐
                if fields and isinstance(fields, list) and any(isinstance(f, dict) for f in fields):
                    for f in fields:
                        if not isinstance(f, dict):
                            continue
                        key = self._get_field_key(f)
                        label = self._get_field_label(f)
                        val = str(job.get(key, '—')) if isinstance(job, dict) else '—'
                        p_field = self.doc.add_paragraph()
                        p_field.style = 'List Bullet'
                        # 所有字段使用相同的缩进，确保项目符号对齐
                        # 不设置额外的left_indent，使用List Bullet的默认缩进
                        run_label = p_field.add_run(f"{label}：")
                        run_label.font.bold = True
                        p_field.add_run(val)
                else:
                    # 兼容旧格式：如果没有fields，使用默认格式
                    title_str = f"{job.get('job_title','—')}（{job.get('company','—')}，{job.get('location','—')}）"
                    p_title = self.doc.add_paragraph()
                    p_title.style = 'List Bullet'
                    run = p_title.add_run(title_str)
                    run.font.bold = True
                    # 处理匹配度和链接（使用相同的对齐）
                    if job.get('match_score'):
                        p_score = self.doc.add_paragraph()
                        p_score.style = 'List Bullet'
                        # 使用List Bullet的默认缩进，确保项目符号对齐
                        p_score.add_run(f"匹配度：{job.get('match_score')}")
                    if job.get('link'):
                        p_link = self.doc.add_paragraph()
                        p_link.style = 'List Bullet'
                        # 使用List Bullet的默认缩进，确保项目符号对齐
                        p_link.add_run(f"链接：{job.get('link')}")
            # self.doc.add_paragraph()
        else:
            # 表格显示：使用 fields 定义列
            if fields:
                headers = [self._get_field_label(f) for f in fields]
            else:
                headers = ['岗位名称', '公司', '地点', '匹配度', '链接']
            
            # 即使无数据也保留表格结构
            display_rows = rows if rows else [{}]
            table = self.doc.add_table(rows=len(display_rows) + 1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头（加粗）
            for i, h in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = h
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.font.bold = True
            
            # 数据行
            for r, job in enumerate(display_rows, start=1):
                if fields:
                    for col_idx, field in enumerate(fields):
                        key = self._get_field_key(field)
                        value = str(job.get(key, '—')) if isinstance(job, dict) else '—'
                        table.cell(r, col_idx).text = value
                else:
                    # 默认字段
                    table.cell(r, 0).text = str(job.get('job_title', '—'))
                    table.cell(r, 1).text = str(job.get('company', '—'))
                    table.cell(r, 2).text = str(job.get('location', '—'))
                    table.cell(r, 3).text = str(job.get('match_score', '—'))
                    table.cell(r, 4).text = str(job.get('link', '—'))
            self.doc.add_paragraph()

    def _add_generic_section(self, title: str, fields: List[Dict[str, Any]], data: Any, add_title: bool = True):
        """未知组件的兜底渲染：标题 + 字段占位或数据"""
        if add_title:
            self._add_section_title(title)
        rows = []
        if isinstance(data, dict) and isinstance(data.get('rows'), list):
            rows = data.get('rows')
        elif isinstance(data, list):
            rows = data
        # 若有多行数据，渲染列表；否则渲染字段列表
        if rows:
            for row in rows:
                p = self.doc.add_paragraph(); p.style = 'List Bullet'
                parts = []
                for f in fields or []:
                    key = f.get('id') or f.get('field') or f.get('name')
                    label = f.get('label') or key
                    val = str(row.get(key, '—')) if isinstance(row, dict) else '—'
                    parts.append(f"{label}:{val}")
                p.add_run("  ".join(parts))
        else:
            # 单次占位
            for f in fields or []:
                if not isinstance(f, dict):
                    continue
                p = self.doc.add_paragraph(); p.style = 'List Bullet'
                label = f.get('label') or f.get('id') or f.get('field') or '字段'
                run = p.add_run(f"{label}："); run.font.bold = True
                p.add_run('—')
        self.doc.add_paragraph()

    def _add_evaluation(self, evaluation: Any):
        """添加评价（文本+要点列表）"""
        if not evaluation:
            return
        self._add_section_title('评价')
        overall = ''
        advantages = []
        risks = []
        advice = ''
        if isinstance(evaluation, dict):
            overall = str(evaluation.get('overall', ''))
            advantages = evaluation.get('advantages', []) or []
            risks = evaluation.get('risks', []) or []
            advice = str(evaluation.get('advice', ''))
        if overall:
            p = self.doc.add_paragraph(overall)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if advantages:
            p = self.doc.add_paragraph(); run = p.add_run('优势：'); run.font.bold = True
            for a in advantages:
                li = self.doc.add_paragraph(); li.style = 'List Bullet'; li.add_run(str(a))
        if risks:
            p = self.doc.add_paragraph(); run = p.add_run('风险：'); run.font.bold = True
            for a in risks:
                li = self.doc.add_paragraph(); li.style = 'List Bullet'; li.add_run(str(a))
        if advice:
            p = self.doc.add_paragraph(); run = p.add_run('建议：'); run.font.bold = True
            self.doc.add_paragraph(advice)
        self.doc.add_paragraph()

    def _add_salary(self, salary: Any):
        """添加薪资（列表）"""
        if not salary:
            return
        self._add_section_title('薪资')
        items = []
        if isinstance(salary, dict):
            items = [
                ('期望薪资', salary.get('expected_salary', '')),
                ('当前薪资', salary.get('current_salary', '')),
                ('目标薪资', salary.get('target_salary', '')),
            ]
        for label, val in items:
            p = self.doc.add_paragraph(); p.style = 'List Bullet'
            run = p.add_run(f'{label}：'); run.font.bold = True
            p.add_run(str(val))
        self.doc.add_paragraph()
    
    def _get_field_key(self, field: Dict[str, Any]) -> str:
        """从字段定义中提取字段key（容错空值/非字典）"""
        if not isinstance(field, dict):
            return ''
        return field.get('id') or field.get('field') or field.get('name') or field.get('key') or ''
    
    def _get_field_label(self, field: Dict[str, Any]) -> str:
        """从字段定义中提取字段标签（容错空值/非字典）"""
        if not isinstance(field, dict):
            return ''
        return field.get('label') or field.get('title') or self._get_field_key(field)
    
    def _add_basic_info_by_fields(self, title: str, fields: List[Dict[str, Any]], sec_data: Any, fallback_data: Dict[str, Any], add_title: bool = True):
        """按字段定义渲染基本信息"""
        if add_title:
            self._add_section_title(title)
        # 模板导出：不使用fallback数据，只使用sec_data或显示占位符
        data = sec_data if isinstance(sec_data, dict) else (fallback_data if fallback_data is not None else {})
        if not fields:
            # 无字段定义时使用默认字段
            default_fields = [
                {'id': 'name', 'label': '姓名'},
                {'id': 'phone', 'label': '电话'},
                {'id': 'email', 'label': '邮箱'},
                {'id': 'location', 'label': '所在地'},
            ]
            fields = default_fields
        
        # 基本信息：字段使用黑点项目符号
        for idx, field in enumerate(fields, 1):
            key = self._get_field_key(field)
            label = self._get_field_label(field)
            value = str(data.get(key, '—')) if isinstance(data, dict) else '—'
            p = self.doc.add_paragraph(); p.style = 'List Bullet'
            run_label = p.add_run(f"{label}：")
            run_label.font.bold = True
            p.add_run(value)
        # self.doc.add_paragraph()
    
    def _add_education_by_fields(self, title: str, config: Dict[str, Any], fields: List[Dict[str, Any]], sec_data: Any, fallback_data: List[Dict[str, Any]], add_title: bool = True):
        """按字段定义和表格列配置渲染教育背景"""
        if add_title:
            self._add_section_title(title)
        rows = []
        if isinstance(sec_data, dict) and isinstance(sec_data.get('rows'), list):
            rows = sec_data.get('rows')
        elif isinstance(sec_data, list):
            rows = sec_data
        else:
            rows = fallback_data if (fallback_data is not None and isinstance(fallback_data, list)) else []
        
        # 使用 config.tableColumns 或根据 fields 生成列
        table_columns = config.get('tableColumns') or []
        if not table_columns and fields:
            # 从 fields 生成列
            safe_fields = [f for f in (fields or []) if isinstance(f, dict)]
            table_columns = [{'id': f.get('id'), 'field': self._get_field_key(f), 'label': self._get_field_label(f)} for f in safe_fields]
        
        if not table_columns:
            # 默认列
            table_columns = [
                {'id': 'period', 'field': 'period', 'label': '起止时间'},
                {'id': 'school', 'field': 'school', 'label': '学校'},
                {'id': 'major', 'field': 'major', 'label': '专业'},
            ]
        
        # 即使无数据也保留表格结构（至少表头+一行占位）
        display_rows = rows if rows else [{}]
        table = self.doc.add_table(rows=len(display_rows) + 1, cols=len(table_columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格外边框为加粗
        self._set_table_outer_border_bold(table)
        
        # 表头（加粗）
        for col_idx, col in enumerate(table_columns):
            col = col if isinstance(col, dict) else {}
            cell = table.cell(0, col_idx)
            cell.text = col.get('label', '')
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 数据行
        for row_idx, row in enumerate(display_rows, 1):
            for col_idx, col in enumerate(table_columns):
                col = col if isinstance(col, dict) else {}
                field_key = col.get('field') or col.get('id')
                value = str(row.get(field_key, '—')) if isinstance(row, dict) else '—'
                table.cell(row_idx, col_idx).text = value
        self.doc.add_paragraph()
    
    def _add_work_experience_by_fields(self, title: str, config: Dict[str, Any], fields: List[Dict[str, Any]], sec_data: Any, fallback_data: List[Dict[str, Any]], add_title: bool = True):
        """按字段定义和表格列配置渲染工作经历，支持 displayMode 配置"""
        if add_title:
            self._add_section_title(title)
        rows = []
        if isinstance(sec_data, dict) and isinstance(sec_data.get('rows'), list):
            rows = sec_data.get('rows')
        elif isinstance(sec_data, list):
            rows = sec_data
        else:
            rows = fallback_data if (fallback_data is not None and isinstance(fallback_data, list)) else []
        
        display_mode = config.get('displayMode', 'summary')  # 'summary' | 'detailed' | 'both'
        details_config = config.get('detailsConfig', {})
        show_map = details_config.get('show', {})
        
        # 渲染表格（summary 或 both 模式）
        if display_mode in ['summary', 'both']:
            table_columns = config.get('tableColumns') or []
            if not table_columns:
                # 默认列
                table_columns = [
                    {'id': 'period', 'field': 'period', 'label': '起止时间'},
                    {'id': 'company', 'field': 'company', 'label': '公司名称'},
                    {'id': 'position', 'field': 'position', 'label': '职位'},
                ]
            
            # 即使无数据也保留表格结构
            display_rows = rows if rows else [{}]
            table = self.doc.add_table(rows=len(display_rows) + 1, cols=len(table_columns))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 设置表格外边框为加粗
            self._set_table_outer_border_bold(table)
            
            # 表头（加粗）
            for col_idx, col in enumerate(table_columns):
                col = col if isinstance(col, dict) else {}
                cell = table.cell(0, col_idx)
                cell.text = col.get('label', '')
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # 数据行
            for row_idx, row in enumerate(display_rows, 1):
                for col_idx, col in enumerate(table_columns):
                    col = col if isinstance(col, dict) else {}
                    field_key = col.get('field') or col.get('id')
                    value = str(row.get(field_key, '—')) if isinstance(row, dict) else '—'
                    cleaned_value = self._clean_text(value) if value != '—' else value
                    table.cell(row_idx, col_idx).text = cleaned_value
            self.doc.add_paragraph()
        
        # 渲染详细卡片（detailed 或 both 模式）
        if display_mode in ['detailed', 'both']:
            if not rows:
                # 即使无数据也保留版式：显示字段占位
                default_keys = ['position', 'report_to', 'team_size', 'location', 'responsibilities', 'achievements', 'reason_for_leaving']
                for key in default_keys:
                    if show_map.get(key, True) is not False:
                        label_map = {
                            'position': '职位',
                            'report_to': '汇报对象',
                            'team_size': '下属团队',
                            'location': '工作地点',
                            'responsibilities': '工作职责',
                            'achievements': '工作业绩',
                            'reason_for_leaving': '离职原因'
                        }
                        p = self.doc.add_paragraph()
                        p.style = 'List Bullet'
                        run_label = p.add_run(f"{label_map.get(key, key)}：")
                        run_label.font.bold = True
                        p.add_run('—')
                self.doc.add_paragraph()
            else:
                for exp in rows:
                    # 标题行：起止时间, 公司名称
                    period = exp.get('period', '')
                    if period:
                        # 如果period已存在，也需要格式化（可能包含 YYYY-MM 格式）
                        # 处理格式如 "2020-01 - 2023-06" 或 "2020.01 - 2023.06"
                        if '-' in period and ' - ' in period:
                            parts = period.split(' - ')
                            period = f"{self._format_date(parts[0].strip())} - {parts[1].strip() if parts[1].strip() != '至今' else '至今'}"
                        elif '-' in period:
                            period = self._format_date(period)
                    else:
                        start_date = self._format_date(exp.get('start_date', ''))
                        end_date = self._format_date(exp.get('end_date', ''))
                        is_current = exp.get('is_current', False)
                        period = f"{start_date} - {'至今' if is_current else end_date}"
                    company = exp.get('company', '')
                    title_line = f"{period}, {company}" if period and company else (period or company or '—')
                    
                    p = self.doc.add_paragraph()
                    run = p.add_run(title_line)
                    run.font.bold = True
                    run.font.underline = True
                    
                    # 默认字段顺序
                    default_keys = [
                        {'key': 'position', 'label': '职位'},
                        {'key': 'report_to', 'label': '汇报对象'},
                        {'key': 'team_size', 'label': '下属团队'},
                        {'key': 'location', 'label': '工作地点'},
                        {'key': 'responsibilities', 'label': '工作职责'},
                        {'key': 'achievements', 'label': '工作业绩'},
                        {'key': 'reason_for_leaving', 'label': '离职原因'},
                    ]
                    
                    for item in default_keys:
                        key = item['key']
                        label = item['label']
                        if show_map.get(key, True) is False:
                            continue
                        value = exp.get(key, '')
                        if not value:
                            continue
                        if key in ['responsibilities', 'achievements', 'reason_for_leaving']:
                            # 数组类型，只显示标签和内容，手动添加编号（每个字段从1开始，字段内连续编号）
                            # 与前端模板组件保持一致：使用HTML的<ol>标签，每个<ol>的编号从1开始
                            if isinstance(value, list) and value:
                                # 添加标签（粗体），使用 List Bullet 样式，与职位等字段保持一致
                                p = self.doc.add_paragraph()
                                p.style = 'List Bullet'
                                run_label = p.add_run(f"{label}：")
                                run_label.font.bold = True
                                # 手动添加编号，每个字段从1开始（与前端<ol>标签行为一致）
                                for idx, v in enumerate(value, 1):
                                    if v and str(v).strip():
                                        # 检查内容是否已经包含编号，如果包含则移除（避免重复编号）
                                        v_str = str(v).strip()
                                        # 移除开头的数字编号格式（如"1. "、"1、"、"1．"、"1)"等）
                                        # 支持阿拉伯数字+常见分隔符（半角/全角点、中文顿号、右括号）
                                        v_str = re.sub(r'^\s*\d+\s*[\.、．\)]\s*', '', v_str)
                                        # 清理特殊字符
                                        v_str = self._clean_text(v_str)
                                        
                                        p_detail = self.doc.add_paragraph()
                                        p_detail.style = 'Normal'
                                        # 设置悬挂缩进：第一行文字可以在编号后有适当空间，第二行及后续行与第一行文字对齐
                                        # left_indent: 控制第二行及后续行的起始位置（与第一行文字对齐）
                                        # first_line_indent: 负值表示悬挂缩进，让第一行向左突出，容纳编号和间距
                                        # 编号"1. "大约占0.3英寸，编号后间距约0.1英寸，所以第一行文字起始位置约0.4英寸
                                        # 第二行及后续行应该与第一行文字对齐，所以left_indent设为0.4英寸
                                        # first_line_indent设为-0.15英寸，让第一行向左突出，容纳编号
                                        p_detail.paragraph_format.left_indent = Inches(0.4)  # 第二行及后续行的起始位置（与第一行文字对齐）
                                        p_detail.paragraph_format.first_line_indent = Inches(-0.15)  # 悬挂缩进，让第一行向左突出以容纳编号
                                        # 设置段后间距6磅
                                        p_detail.paragraph_format.space_after = Pt(6)
                                        # 手动添加编号文本（"1. "、"2. "等）
                                        p_detail.add_run(f"{idx}. {v_str}")
                        else:
                            # 普通字段
                            p = self.doc.add_paragraph()
                            p.style = 'List Bullet'
                            run_label = p.add_run(f"{label}：")
                            run_label.font.bold = True
                            cleaned_value = self._clean_text(str(value)) if value else ''
                            p.add_run(cleaned_value)
                    
                    self.doc.add_paragraph()  # 每个经历之间空一行
    
    def _add_evaluation_by_fields(self, title: str, fields: List[Dict[str, Any]], sec_data: Any, fallback_data: Dict[str, Any], add_title: bool = True):
        """按字段定义渲染评价"""
        if add_title:
            self._add_section_title(title)
        data = sec_data if isinstance(sec_data, dict) else (fallback_data if (fallback_data is not None and isinstance(fallback_data, dict)) else {})
        if not isinstance(data, dict):
            data = {}
        
        if not fields:
            # 默认字段
            fields = [
                {'id': 'overall', 'label': '总体评价'},
                {'id': 'advantages', 'label': '优势'},
                {'id': 'risks', 'label': '风险'},
                {'id': 'advice', 'label': '建议'},
            ]
        
        for field in fields:
            key = self._get_field_key(field)
            label = self._get_field_label(field)
            value = data.get(key, '')
            
            # 空数据处理：采用与工作经历和项目经历相同的处理方式，不显示字段
            if not value:
                continue
            
            if isinstance(value, list):
                # 列表类型（如优势、风险）
                if not value:  # 空列表也不显示
                    continue
                p = self.doc.add_paragraph()
                p.style = 'List Bullet'
                run_label = p.add_run(f"{label}：")
                run_label.font.bold = True
                for item in value:
                    p_item = self.doc.add_paragraph()
                    p_item.paragraph_format.left_indent = Inches(0.5)
                    p_item.add_run(f"  • {str(item)}")
            else:
                # 文本类型：采用与工作职责相同的格式（编号列表）
                text_value = str(value) if value else ''
                if not text_value or not text_value.strip():
                    # 空数据不显示字段（与工作经历和项目经历保持一致）
                    continue
                
                # 添加标签（粗体），使用 List Bullet 样式，与工作职责保持一致
                p = self.doc.add_paragraph()
                p.style = 'List Bullet'
                run_label = p.add_run(f"{label}：")
                run_label.font.bold = True
                
                # 处理换行符：将双换行(\n\n)转换为段落分隔，单个换行(\n)替换为空格
                cleaned_value = self._clean_text(text_value)
                
                # 智能分割：支持多种分隔符来分割条目
                # 1. 优先检查是否已有编号格式（如"1. "、"2. "、"1、"等）
                numbered_pattern = r'\n\s*\d+\s*[\.、．\)]\s+'
                if re.search(numbered_pattern, cleaned_value):
                    # 按编号分割
                    paragraphs = re.split(numbered_pattern, cleaned_value)
                    paragraphs = [para.strip() for para in paragraphs if para.strip()]
                # 2. 检查双换行（段落分隔）
                elif '\n\n' in cleaned_value:
                    paragraphs = [para.strip() for para in cleaned_value.split('\n\n') if para.strip()]
                # 3. 检查单个换行（如果文本较长，可能是条目分隔）
                elif '\n' in cleaned_value and len(cleaned_value) > 100:
                    # 对于较长的文本，单个换行也可能表示条目分隔
                    paragraphs = [para.strip() for para in cleaned_value.split('\n') if para.strip()]
                    # 如果分割后条目太多（可能是误分割），合并回去
                    if len(paragraphs) > 10:
                        paragraphs = [cleaned_value]
                # 4. 检查分号分隔（中文分号或英文分号）
                elif '；' in cleaned_value or ';' in cleaned_value:
                    # 按分号分割，但需要过滤掉太短的片段（可能是误分割）
                    parts = re.split(r'[；;]', cleaned_value)
                    paragraphs = [para.strip() for para in parts if para.strip() and len(para.strip()) > 10]
                    # 如果分割后条目太少，可能不是分号分隔，保持原样
                    if len(paragraphs) < 2:
                        paragraphs = [cleaned_value] if cleaned_value.strip() else []
                else:
                    # 没有明显的分隔符，整个文本作为一个条目
                    paragraphs = [cleaned_value] if cleaned_value.strip() else []
                
                # 手动添加编号，每个字段从1开始（与工作职责格式一致）
                for idx, para in enumerate(paragraphs, 1):
                    if not para or not para.strip():
                        continue
                    
                    # 检查内容是否已经包含编号，如果包含则移除（避免重复编号）
                    para_str = para.strip()
                    # 移除开头的数字编号格式（如"1. "、"1、"、"1．"、"1)"等）
                    para_str = re.sub(r'^\s*\d+\s*[\.、．\)]\s*', '', para_str)
                    # 清理特殊字符（已在_clean_text中处理，但确保没有残留编号）
                    para_str = self._clean_text(para_str)
                    
                    # 创建编号条目，使用与工作职责相同的格式
                    li = self.doc.add_paragraph()
                    li.style = 'Normal'
                    # 设置悬挂缩进：与工作职责保持一致
                    # left_indent: 控制第二行及后续行的起始位置（与第一行文字对齐）
                    # first_line_indent: 负值表示悬挂缩进，让第一行向左突出，容纳编号和间距
                    li.paragraph_format.left_indent = Inches(0.4)  # 第二行及后续行的起始位置（与第一行文字对齐）
                    li.paragraph_format.first_line_indent = Inches(-0.15)  # 悬挂缩进，让第一行向左突出以容纳编号
                    # 手动添加编号文本（"1. "、"2. "等）
                    li.add_run(f"{idx}. {para_str}")
        self.doc.add_paragraph()
    
    def _add_salary_by_fields(self, title: str, fields: List[Dict[str, Any]], sec_data: Any, fallback_data: Dict[str, Any], add_title: bool = True):
        """按字段定义渲染薪资"""
        if add_title:
            self._add_section_title(title)
        data = sec_data if isinstance(sec_data, dict) else (fallback_data if (fallback_data is not None and isinstance(fallback_data, dict)) else {})
        if not isinstance(data, dict):
            data = {}
        
        if not fields:
            # 默认字段
            fields = [
                {'id': 'expected_salary', 'label': '期望薪资'},
                {'id': 'current_salary', 'label': '当前薪资'},
                {'id': 'target_salary', 'label': '目标薪资'},
            ]
        
        for field in fields:
            key = self._get_field_key(field)
            label = self._get_field_label(field)
            value = str(data.get(key, '—'))
            p = self.doc.add_paragraph()
            p.style = 'List Bullet'
            run_label = p.add_run(f"{label}：")
            run_label.font.bold = True
            p.add_run(value)
        self.doc.add_paragraph()
    
    def _add_skills_by_fields(self, title: str, fields: List[Dict[str, Any]], sec_data: Any, fallback_data: Dict[str, Any], add_title: bool = True):
        """按字段定义渲染技能"""
        if add_title:
            self._add_section_title(title)
        data = sec_data if isinstance(sec_data, dict) else (fallback_data if (fallback_data is not None and isinstance(fallback_data, dict)) else {})
        if not isinstance(data, dict):
            data = {}
        
        if not fields:
            # 默认字段
            fields = [
                {'id': 'technical', 'label': '技术技能'},
                {'id': 'soft', 'label': '软技能'},
                {'id': 'languages', 'label': '语言能力'},
            ]
        
        for field in fields:
            key = self._get_field_key(field)
            label = self._get_field_label(field)
            value = data.get(key, '')
            
            p = self.doc.add_paragraph()
            run_label = p.add_run(f"{label}：")
            run_label.font.bold = True
            
            if isinstance(value, list):
                p.add_run(", ".join([str(v) for v in value]) if value else '—')
            else:
                p.add_run(str(value) if value else '—')
        self.doc.add_paragraph()
    
    def _add_projects_by_fields(self, title: str, fields: List[Dict[str, Any]], sec_data: Any, fallback_data: List[Dict[str, Any]], add_title: bool = True):
        """按字段定义渲染项目经历"""
        if add_title:
            self._add_section_title(title)
        rows = []
        if isinstance(sec_data, dict) and isinstance(sec_data.get('rows'), list):
            rows = sec_data.get('rows')
        elif isinstance(sec_data, list):
            rows = sec_data
        else:
            rows = fallback_data if (fallback_data is not None and isinstance(fallback_data, list)) else []

        # 如果没有数据，不渲染（整个组件不显示）
        if not rows or len(rows) == 0:
            return

        # 如果有后备 projects 数据，尝试按索引或项目名称匹配，为每一行补充 start_date/end_date
        if isinstance(fallback_data, list) and rows:
            merged_rows = []
            for idx, row in enumerate(rows):
                base_row = row if isinstance(row, dict) else {}
                # 优先按索引匹配，如果索引超出范围或项目名称不匹配，尝试按名称匹配
                proj_src = None
                if idx < len(fallback_data) and isinstance(fallback_data[idx], dict):
                    proj_src = fallback_data[idx]
                    # 如果项目名称不匹配，尝试按名称查找
                    row_name = base_row.get('name') or base_row.get('project_name') or ''
                    src_name = proj_src.get('name') or proj_src.get('project_name') or ''
                    if row_name and src_name and row_name != src_name:
                        # 按名称查找匹配的项目
                        for proj in fallback_data:
                            if isinstance(proj, dict):
                                proj_name = proj.get('name') or proj.get('project_name') or ''
                                if proj_name == row_name:
                                    proj_src = proj
                                    break
                
                # 如果行中没有对应字段或字段值为空，从后备数据补充
                if proj_src:
                    if not base_row.get('start_date') and proj_src.get('start_date'):
                        base_row['start_date'] = proj_src.get('start_date')
                    if not base_row.get('end_date') and proj_src.get('end_date'):
                        base_row['end_date'] = proj_src.get('end_date')
                merged_rows.append(base_row)
            rows = merged_rows
        
        if not fields:
            # 默认字段
            fields = [
                {'id': 'name', 'label': '项目名称'},
                {'id': 'role', 'label': '角色'},
                {'id': 'description', 'label': '描述'},
                {'id': 'outcome', 'label': '成果'},
            ]
        
        # 渲染项目数据
        display_rows = rows
        for row_idx, row in enumerate(display_rows):
            # 计算该项目的时间段，用于附加到名称后
            start_raw = row.get('start_date') if isinstance(row, dict) else None
            end_raw = row.get('end_date') if isinstance(row, dict) else None
            # 转换为字符串并去除空白，空字符串转为 None
            if start_raw:
                start_raw = str(start_raw).strip() or None
            if end_raw:
                end_raw = str(end_raw).strip() or None
            start_fmt = self._format_date(start_raw) if start_raw else ''
            end_fmt = self._format_date(end_raw) if end_raw else ''
            period = ''
            if start_fmt and end_fmt:
                period = f"{start_fmt} - {end_fmt}"
            elif start_fmt and not end_fmt:
                period = f"{start_fmt} - 至今"
            elif not start_fmt and end_fmt:
                period = end_fmt

            for field in fields:
                key = self._get_field_key(field)
                label = self._get_field_label(field)
                raw_value = row.get(key, '') if isinstance(row, dict) else ''
                
                # 处理值：如果是数组，合并成字符串；如果是字符串，直接使用
                if isinstance(raw_value, list):
                    # 如果是数组，合并成字符串（用空格连接，避免引入换行）
                    value = ' '.join([str(v).strip() for v in raw_value if v and str(v).strip()]) if raw_value else ''
                elif isinstance(raw_value, str):
                    value = raw_value.strip() if raw_value else ''
                else:
                    value = str(raw_value).strip() if raw_value else ''

                # 如果当前字段是项目名称，并且有时间段，追加到名称后
                # 支持多种可能的字段名：name, project_name, 项目名称等
                name_keys = ['name', 'project_name', '项目名称']
                if key in name_keys and period:
                    # 避免重复添加括号（如果名称中已经包含时间段，不再追加）
                    if value:
                        # 检查是否已经包含时间段格式（如 "2020.02 - 至今"）
                        if ' - ' not in value or '至今' not in value:
                            value = f"{value} ({period})"
                    else:
                        value = period
                
                # 如果值为空，不显示该字段（与工作经历的处理方式一致）
                if not value:
                    continue
                
                # 第二个及后续项目的项目名称字段：先添加空段落（换行），确保段前间距生效
                # 在确认值不为空后再添加，避免空字段也添加换行
                if row_idx > 0 and key in name_keys:
                    self.doc.add_paragraph()
                
                # 清理特殊字符
                value = self._clean_text(value)
                
                p = self.doc.add_paragraph()
                p.style = 'List Bullet'
                run_label = p.add_run(f"{label}：")
                run_label.font.bold = True
                p.add_run(value)
                # 设置段后间距6磅
                p.paragraph_format.space_after = Pt(6)
            # self.doc.add_paragraph()
    
    def _add_section_title(self, title: str):
        """添加章节标题"""
        heading = self.doc.add_heading(level=2)
        heading_run = heading.add_run(title)
        heading_run.font.size = Pt(12)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # 添加分割线
        # self.doc.add_paragraph("_" * 50)
    
    def _set_table_outer_border_bold(self, table):
        """
        设置表格外边框为加粗（2.25pt，约等于 Word 中的"粗线"）
        内边框保持默认（0.5pt）
        注意：Word 中边框宽度单位是 1/8 point，所以 2.25pt = 18/8 point = 18
        """
        # 外边框宽度：2.25pt（粗线）= 18/8 point
        outer_border_width = "18"
        
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                
                # 创建或获取边框元素
                tcBorders = OxmlElement('w:tcBorders')
                
                # 判断是否为外边框
                is_top = row_idx == 0
                is_bottom = row_idx == num_rows - 1
                is_left = col_idx == 0
                is_right = col_idx == num_cols - 1
                
                # 设置边框
                borders = {}
                if is_top:
                    borders['top'] = outer_border_width
                else:
                    borders['top'] = "4"  # 默认内边框 0.5pt = 4/8 point
                
                if is_bottom:
                    borders['bottom'] = outer_border_width
                else:
                    borders['bottom'] = "4"
                
                if is_left:
                    borders['left'] = outer_border_width
                else:
                    borders['left'] = "4"
                
                if is_right:
                    borders['right'] = outer_border_width
                else:
                    borders['right'] = "4"
                
                # 创建边框元素
                for border_name, width in borders.items():
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), width)
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                # 移除旧的边框元素（如果存在）
                old_borders = tcPr.find(qn('w:tcBorders'))
                if old_borders is not None:
                    tcPr.remove(old_borders)
                
                # 添加新的边框元素
                tcPr.append(tcBorders)
    
    def _generate_filename(self, resume_data: Dict[str, Any]) -> str:
        """生成文件名（清理非法字符，优先使用人名）"""
        import re
        # 文件名生成：优先 basic_info.name（人名）；其次 template_name；否则 'resume'
        name = 'resume'
        if isinstance(resume_data, dict):
            # 优先从 basic_info 中获取人名
            bi = resume_data.get('basic_info')
            if isinstance(bi, dict):
                bn = bi.get('name')
                if isinstance(bn, str) and bn.strip():
                    name = bn.strip()
            
            # 如果没有人名，再尝试使用 template_name
            if name == 'resume':
                tmpl_name = resume_data.get('template_name')
                if isinstance(tmpl_name, str) and tmpl_name.strip():
                    name = tmpl_name.strip()
        
        # 清理文件名中的非法字符（Windows/Linux都不允许的字符）
        # 移除或替换：/ \ : * ? " < > |
        illegal_chars = r'[/\\:*?"<>|]'
        name = re.sub(illegal_chars, '_', name)
        # 移除多余的下划线和空格
        name = re.sub(r'_{2,}', '_', name).strip('_').strip()
        # 限制文件名长度（避免过长）
        if len(name) > 50:
            name = name[:50]
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # 增加唯一后缀，避免并发覆盖
        try:
            import uuid
            unique = uuid.uuid4().hex[:8]
        except Exception:
            unique = "00000000"
        filename = f"{name}_简历_{timestamp}_{unique}.docx"
        
        # 保存文档
        export_dir = settings.export_dir or 'temp'
        save_path = os.path.join(export_dir, filename)
        os.makedirs(export_dir, exist_ok=True)
        self.doc.save(save_path)
        
        return save_path
    

# 全局导出器实例
word_exporter = WordExporter()